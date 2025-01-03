import os
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Inches
import streamlit as st
import pythoncom
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter

# Configure page settings
st.set_page_config(
    page_title="Affidavit Generator",
    page_icon="📝"
)


# Function to load CSS
def local_css(file_name):
    with open(file_name) as css:
        st.markdown("<style>{}</style>".format(css.read()), unsafe_allow_html=True)

# Call function
local_css("styles/main.css")


# Path to template files
affidavit_template_path = "Affidavit template to change.docx"
exhibit_template_path = "Exhibit Template.docx"
part_2_template_path = "Affidavit template to change - Part 2.docx"

# Helper function to handle DOCX to PDF conversion
def convert_with_com(input_path, output_path):
    pythoncom.CoInitialize()  # Initialize COM for the current thread
    try:
        convert(input_path, output_path)
    finally:
        pythoncom.CoUninitialize()  # Uninitialize COM after use

# Streamlit form for user inputs
st.title("Affidavit Generator with Exhibits")

name = st.text_input("Enter the name:").upper()
case_file = st.text_input("What is your case file number?")
party_name = st.text_input("What is your party name:")
lawyer_name = st.text_input("What is your lawyer's name?")
date = st.date_input("What date do you want?")
stat_declaration = st.selectbox("What is your statutory declaration?", ["Sworn", "Affirmed"])
address = st.text_area("What is your address?")
email = st.text_input("What is your email?")
phone = st.text_input("What is your phone number?")
party_role = st.selectbox("Enter your role in this proceeding (e.g., witness, plaintiff, defendant):", ["Witness", "Plaintiff", "Defendant"])

# Streamlit file uploader for exhibits
exhibit_files = st.file_uploader("Upload Exhibits", type=["pdf"], accept_multiple_files=True)

# Generate affidavit and exhibits
def generate_affidavit_with_exhibits(output_filename):
    # Prepare the context for affidavit rendering
    context = {
        'name': name,
        'case_file': case_file,
        'party_name': party_name,
        'date': date.strftime('%Y-%m-%d'),  # Format date for consistency
        'stat_declaration': stat_declaration,
        'address': address,
        'email': email,
        'phone': phone,
        'lawyer_name': lawyer_name,
        'name': name,
        'party_role': party_role.lower()
    }

    # Load the affidavit template
    doc = DocxTemplate(affidavit_template_path)
    doc.render(context)

    # Save the rendered affidavit into a temporary file
    affidavit_file_path = "temp_affidavit.docx"
    doc.save(affidavit_file_path)

    # Load the affidavit document again using python-docx
    affidavit_doc = Document(affidavit_file_path)

    # Process each exhibit file uploaded by the user
    exhibit_letter = 'a'  # Start with Exhibit A

    for exhibit_file in exhibit_files:
        # Load the exhibit template as a DocxTemplate
        exhibit_doc = DocxTemplate(exhibit_template_path)
        
        # Prepare context for the exhibit
        exhibit_context = {
            'letter': exhibit_letter,
            'party_name': party_name,
            'date': date.strftime('%Y-%m-%d')
        }
        
        # Render the exhibit template with the context
        exhibit_doc.render(exhibit_context)
        
        # Save rendered exhibit to a temporary file
        temp_exhibit_path = f"temp_exhibit_{exhibit_letter}.docx"
        exhibit_doc.save(temp_exhibit_path)
        
        # Append the rendered exhibit to the affidavit document
        rendered_exhibit = Document(temp_exhibit_path)
        for element in rendered_exhibit.element.body:
            affidavit_doc.element.body.append(element)
        
        # Increment the exhibit letter for the next exhibit
        exhibit_letter = chr(ord(exhibit_letter) + 1)

    # Append the part 2 document
    part_2_doc = DocxTemplate(part_2_template_path)

    # Prepare context for the part 2 document
    part_2_context = {
        'party_name': party_name,
        'party_role': party_role,
        'lawyer_name': lawyer_name,
        'date': date.strftime('%Y-%m-%d'),
        'name': name,
        'address': address,
        'email': email,
        'phone': phone,
        'stat_declaration': stat_declaration,
        'case_file': case_file
    }

    # Render the part 2 template with the context
    part_2_doc.render(part_2_context)

    # Save rendered part 2 to a temporary file
    temp_part_2_path = "temp_part_2.docx"
    part_2_doc.save(temp_part_2_path)

    # Append the rendered part 2 to the affidavit document
    rendered_part_2 = Document(temp_part_2_path)
    for element in rendered_part_2.element.body:
        affidavit_doc.element.body.append(element)

    # Apply 1-inch margins to the final document
    for section in affidavit_doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Save the final document with affidavit, exhibits, and part 2 of the affidavit
    affidavit_doc.save(output_filename)

# Word file generation
if st.button("Generate Affidavit with Exhibits (Word)"):
    output_filename = f"{name}_affidavit_with_exhibits.docx"
    generate_affidavit_with_exhibits(output_filename)
    with open(output_filename, "rb") as f:
        st.download_button(
            label="Download Affidavit with Exhibits (Word)",
            data=f,
            file_name=output_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# PDF file generation
uploaded_word_file = st.file_uploader("Upload the Word file to convert to PDF", type=["docx"])

if uploaded_word_file:
    st.subheader("PDF Generation")

    def generate_pdf_with_exhibits():
        # Save uploaded Word file temporarily
        temp_word_path = "uploaded_temp_affidavit.docx"
        with open(temp_word_path, "wb") as f:
            f.write(uploaded_word_file.getbuffer())

        # Convert Word to PDF
        pdf_output_path = temp_word_path.replace(".docx", "_converted.pdf")
        convert_with_com(temp_word_path, pdf_output_path)

        # Read the converted affidavit PDF
        affidavit_pdf = PdfReader(pdf_output_path)

        # Initialize a new PDF writer
        pdf_writer = PdfWriter()

        # Add first part of the affidavit template
        affidavit_part_1_docx = "Affidavit template to change.docx"
        affidavit_part_1_template = DocxTemplate(affidavit_part_1_docx)
        affidavit_part_1_context = {
            'name': name,
            'case_file': case_file,
            'party_name': party_name,
            'date': date.strftime('%Y-%m-%d'),  # Format date for consistency
            'stat_declaration': stat_declaration,
            'address': address,
            'email': email,
            'phone': phone,
            'party_role': party_role.lower(),
            'lawyer_name': lawyer_name 
        }

        affidavit_part_1_template.render(affidavit_part_1_context)
        affidavit_part_1_path = "temp_affidavit_part_1.docx"
        affidavit_part_1_template.save(affidavit_part_1_path)
        convert_with_com(affidavit_part_1_path, "temp_affidavit_part_1.pdf")
        
        part_1_pdf = PdfReader("temp_affidavit_part_1.pdf")
        for page in part_1_pdf.pages:
            pdf_writer.add_page(page)

        # Insert exhibit templates and exhibits in the specified order
        exhibit_letter = 'a'
        for exhibit_file in exhibit_files:
            # Convert Exhibit Template to PDF with current exhibit letter
            exhibit_context = {
                'letter': exhibit_letter.upper(),
                'party_name': party_name,
                'date': date.strftime('%Y-%m-%d')
            }
            temp_exhibit_docx = f"temp_exhibit_{exhibit_letter}.docx"
            temp_exhibit_pdf = temp_exhibit_docx.replace(".docx", ".pdf")
            
            exhibit_template = DocxTemplate(exhibit_template_path)
            exhibit_template.render(exhibit_context)
            exhibit_template.save(temp_exhibit_docx)
            convert_with_com(temp_exhibit_docx, temp_exhibit_pdf)

            # Add the exhibit template PDF
            template_pdf = PdfReader(temp_exhibit_pdf)
            for page in template_pdf.pages:
                pdf_writer.add_page(page)

            # Add the uploaded exhibit PDF
            exhibit_pdf = PdfReader(exhibit_file)
            for page in exhibit_pdf.pages:
                pdf_writer.add_page(page)

            # Increment exhibit letter
            exhibit_letter = chr(ord(exhibit_letter) + 1)

        # Convert Part 2 Template to PDF and append
        temp_part_2_docx = "temp_part_2.docx"
        temp_part_2_pdf = temp_part_2_docx.replace(".docx", ".pdf")

        part_2_context = {
            'party_name': party_name,
            'party_role': party_role,
            'lawyer_name': lawyer_name,  # Add this line
            'date': date.strftime('%Y-%m-%d'),
            'name': name,
            'address': address,
            'email': email,
            'phone': phone,
            'stat_declaration': stat_declaration,
            'case_file': case_file
        }

        part_2_template = DocxTemplate(part_2_template_path)
        part_2_template.render(part_2_context)
        part_2_template.save(temp_part_2_docx)
        convert_with_com(temp_part_2_docx, temp_part_2_pdf)

        part_2_pdf = PdfReader(temp_part_2_pdf)
        for page in part_2_pdf.pages:
            pdf_writer.add_page(page)

        # Write the final PDF
        final_pdf_output_path = "affidavit_with_exhibits.pdf"
        with open(final_pdf_output_path, "wb") as output_pdf:
            pdf_writer.write(output_pdf)

        # Provide download link
        with open(final_pdf_output_path, "rb") as f:
            st.download_button(
                label="Download Affidavit with Exhibits (PDF)",
                data=f,
                file_name=os.path.basename(final_pdf_output_path),
                mime="application/pdf"
            )

    generate_pdf_with_exhibits()
